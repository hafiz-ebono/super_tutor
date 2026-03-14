"""
Memory-only document text extraction for PDF and DOCX files.

Never writes to the filesystem. Uses BytesIO throughout.
Raises DocumentExtractionError for unrecoverable conditions.

Phase 12 wraps calls to extract_document() with asyncio.to_thread()
since pypdf and python-docx are blocking synchronous libraries.
"""
from io import BytesIO
import logging

from pypdf import PdfReader
from docx import Document

from app.extraction.cleaner import clean_extracted_content
from app.config import get_settings

logger = logging.getLogger("super_tutor.extraction")

# Module-level constants expose the default values for introspection and testing.
# Runtime code reads the live config via get_settings() so these can be overridden
# without restarting by changing the env vars (honoured on next Settings instantiation).
SCANNED_PDF_THRESHOLD: int = 200
TRUNCATION_LIMIT: int = 50_000
TRUNCATION_MARKER: str = (
    f"\n\n[Content truncated: document exceeds {TRUNCATION_LIMIT:,} characters. "
    "Upload a specific chapter or section for full coverage.]"
)


class DocumentExtractionError(Exception):
    """Raised for unrecoverable document extraction failures."""

    def __init__(self, error_kind: str, message: str = ""):
        self.error_kind = error_kind
        self.message = message
        super().__init__(message or error_kind)


def extract_document(data: bytes, filename: str) -> tuple[str, bool]:
    """
    Extract plain text from PDF or DOCX bytes.

    Args:
        data: Raw file bytes.
        filename: Original filename — used for extension-based dispatch.

    Returns:
        Tuple of (cleaned_text, was_truncated). was_truncated is True when the
        document exceeded the configured character limit and was cut short.

    Raises:
        DocumentExtractionError: For scanned PDFs, unsupported formats.
    """
    fname = filename.lower()
    if fname.endswith(".pdf"):
        raw = _extract_pdf(data)
    elif fname.endswith(".docx"):
        raw = _extract_docx(data)
    else:
        raise DocumentExtractionError(
            error_kind="unsupported_format",
            message=f"Unsupported file type: {filename}. Supported: PDF, DOCX.",
        )

    # Soft truncation before cleaning (truncation marker must survive cleaning)
    truncated, was_truncated = _soft_truncate(raw)

    return clean_extracted_content(truncated, source_type="document"), was_truncated


def _extract_pdf(pdf_bytes: bytes) -> str:
    """Extract text from PDF bytes using pypdf. Raises for scanned/image PDFs."""
    reader = PdfReader(BytesIO(pdf_bytes))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        parts.append(text)
    raw = "\n\n".join(parts)

    if len(raw.strip()) < get_settings().scanned_pdf_threshold:
        raise DocumentExtractionError(
            error_kind="scanned_pdf",
            message=(
                "This PDF appears to be scanned or image-based. "
                "No readable text could be extracted."
            ),
        )
    return raw


def _extract_docx(docx_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx.

    Extracts non-empty paragraphs and table cells.
    Note: merged cells in complex tables may produce duplicate text;
    deduplicate by cell identity if this becomes an issue in production.
    """
    doc = Document(BytesIO(docx_bytes))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n\n".join(parts)


def _soft_truncate(text: str) -> tuple[str, bool]:
    """Truncate text at the nearest paragraph or sentence boundary below the configured limit.

    Returns:
        Tuple of (text, was_truncated). was_truncated is True only when the text
        exceeded the limit and was cut.
    """
    limit = get_settings().document_truncation_limit
    if len(text) <= limit:
        return text, False

    # Prefer paragraph boundary
    boundary = text.rfind("\n\n", 0, limit)
    if boundary == -1:
        # Fall back to sentence boundary
        boundary = text.rfind(". ", 0, limit)
    if boundary == -1:
        # Hard cut as last resort
        boundary = limit

    logger.warning(
        "Document truncated at char %d (original length: %d)",
        boundary,
        len(text),
    )
    marker = (
        f"\n\n[Content truncated: document exceeds {limit:,} characters. "
        "Upload a specific chapter or section for full coverage.]"
    )
    return text[:boundary] + marker, True
